"""
Extração e limpeza de texto do corpus (Etapa 3 do desafio).

Princípios:
- Extrai texto de PDF / DOCX / HTML / TXT.
- Remove ruído típico de PDFs técnicos: cabeçalhos/rodapés repetidos,
  números de página e a seção de Referências bibliográficas.
- Normaliza encoding (mojibake), Unicode, hifenização de quebra de linha e
  espaços múltiplos.
- PRESERVA tabelas de parâmetros e requisitos normativos numerados: blocos
  "estruturados" (listas, artigos, créditos, linhas de tabela) mantêm suas
  quebras de linha, enquanto blocos de prosa têm as quebras de linha de
  diagramação unidas em parágrafos.

Uso típico:
    from src.textproc import extract_and_clean
    res = extract_and_clean("corpus/raw/manual.pdf")
    texto = res["text"]
"""
from __future__ import annotations

import unicodedata
from collections import Counter
from pathlib import Path

try:
    import regex as re  # suporta \p{...} (Unicode property)
except ImportError:  # pragma: no cover
    import re  # type: ignore

# Extração bruta por formato

def _extract_pdf(path: Path) -> list[str]:
    """Retorna o texto de cada página (lista de strings) usando pdfplumber."""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
            pages.append(txt)
    return pages


def _extract_docx(path: Path) -> list[str]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    # Tabelas: preservadas como linhas separadas por " | "
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return ["\n".join(parts)]


def _extract_html(path: Path) -> list[str]:
    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    return [soup.get_text("\n")]


def _extract_txt(path: Path) -> list[str]:
    return [path.read_text(encoding="utf-8", errors="ignore")]


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".txt": _extract_txt,
}


# Remoção de cabeçalhos / rodapés repetidos (apenas multi-página, ex.: PDF)

def _norm_key(line: str) -> str:
    """Normaliza uma linha para comparar repetição (ignora dígitos)."""
    s = re.sub(r"\d+", "#", line).strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s


def detect_repeated_lines(pages: list[str], zone: int = 3, min_ratio: float = 0.5) -> set[str]:
    """Detecta linhas que se repetem no topo/rodapé da maioria das páginas."""
    if len(pages) < 4:
        return set()
    counter: Counter[str] = Counter()
    for pg in pages:
        lines = [ln for ln in pg.splitlines() if ln.strip()]
        zone_lines = lines[:zone] + lines[-zone:]
        for ln in set(_norm_key(ln) for ln in zone_lines):
            if ln:
                counter[ln] += 1
    threshold = max(3, int(len(pages) * min_ratio))
    return {key for key, c in counter.items() if c >= threshold}


_PAGE_NUM_RE = re.compile(
    r"^\s*(p[áa]g(?:ina)?\.?\s*)?\d{1,4}\s*(/|de|of)?\s*\d{0,4}\s*$",
    re.IGNORECASE,
)


def _strip_headers_footers(pages: list[str], repeated: set[str]) -> str:
    out: list[str] = []
    for pg in pages:
        kept: list[str] = []
        for ln in pg.splitlines():
            if _norm_key(ln) in repeated:
                continue
            if _PAGE_NUM_RE.match(ln):
                continue
            kept.append(ln)
        out.append("\n".join(kept))
    return "\n\n".join(out)


# Remoção da seção de Referências (conservadora: só corta um bloco final óbvio)

_REF_HEADING_RE = re.compile(
    r"^\s*(refer[êe]ncias(\s+bibliogr[áa]ficas)?|bibliografia|references)\s*:?\s*$",
    re.IGNORECASE,
)


def _strip_references_section(text: str) -> str:
    """Remove tudo a partir de um título 'Referências/Bibliografia' que esteja
    no terço final do documento (evita cortar requisitos numerados no meio)."""
    lines = text.splitlines()
    n = len(lines)
    for i, ln in enumerate(lines):
        if i > n * 0.6 and _REF_HEADING_RE.match(ln):
            return "\n".join(lines[:i])
    return text


# Normalização e reconstrução de parágrafos

def _fix_encoding(text: str) -> str:
    try:
        import ftfy
        text = ftfy.fix_text(text)
    except ImportError:
        pass
    text = unicodedata.normalize("NFKC", text)
    # espaços/zero-width exóticos -> espaço normal
    text = re.sub(r"[  -​  　]", " ", text)
    text = text.replace("﻿", "")
    return text


def _dehyphenate(text: str) -> str:
    """Junta palavras quebradas por hífen no fim da linha (ex.: 'efici-\\nência')."""
    return re.sub(r"(\p{Ll})-\n(\p{Ll})", r"\1\2", text)


# Marcadores de bloco "estruturado" que NÃO devem ter as linhas unidas
_STRUCT_PREFIX = re.compile(
    r"^\s*("
    r"\d+(\.\d+)*[\).\-–]\s"           # 1.  1.2.3)  4-
    r"|[a-zA-Z][\).]\s"                 # a)  b.
    r"|[•\-–*]\s"                       # bullets
    r"|art\.?\s*\d+"                    # Art. 12
    r"|cr[ée]dito\s"                    # Crédito ...
    r"|tabela\s*\d+"                    # Tabela 3
    r"|requisito\s"                     # Requisito ...
    r")",
    re.IGNORECASE,
)


def _looks_structured(block: str) -> bool:
    """Heurística: tabela / lista / requisito numerado -> preservar quebras."""
    lines = [ln for ln in block.splitlines() if ln.strip()]
    if len(lines) <= 1:
        return False
    # muitas linhas curtas, ou com prefixos de enumeração, ou separador de coluna
    short = sum(1 for ln in lines if len(ln.strip()) < 50)
    enum = sum(1 for ln in lines if _STRUCT_PREFIX.match(ln))
    pipes = sum(1 for ln in lines if ("|" in ln or "\t" in ln))
    digits = sum(1 for ln in lines if re.search(r"\d", ln))
    if pipes >= max(2, len(lines) * 0.5):
        return True
    if enum >= max(2, len(lines) * 0.4):
        return True
    if short >= len(lines) * 0.7 and digits >= len(lines) * 0.5:
        return True
    return False


def _join_prose(block: str) -> str:
    """Une linhas de diagramação de um parágrafo de prosa em uma só linha."""
    block = re.sub(r"[ \t]*\n[ \t]*", " ", block)
    block = re.sub(r"[ \t]{2,}", " ", block)
    return block.strip()


def _normalize_blocks(text: str) -> str:
    # quebra em blocos por linha em branco
    raw_blocks = re.split(r"\n\s*\n", text)
    out_blocks: list[str] = []
    for block in raw_blocks:
        block = block.strip("\n")
        if not block.strip():
            continue
        if _looks_structured(block):
            # preserva quebras; só limpa espaços horizontais excessivos
            lines = [re.sub(r"[ \t]{2,}", " ", ln.rstrip()) for ln in block.splitlines()]
            out_blocks.append("\n".join(ln for ln in lines if ln.strip()))
        else:
            out_blocks.append(_join_prose(block))
    return "\n\n".join(out_blocks)


# API pública

def clean_text(text: str, *, drop_references: bool = True) -> str:
    text = _fix_encoding(text)
    text = _dehyphenate(text)
    if drop_references:
        text = _strip_references_section(text)
    text = _normalize_blocks(text)
    # colapsa 3+ linhas em branco em uma separação de parágrafo
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_and_clean(path: str | Path, *, drop_references: bool = True) -> dict:
    """Extrai e limpa um documento. Retorna dict com texto e diagnóstico."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in _EXTRACTORS:
        raise ValueError(f"Formato não suportado: {ext} ({path.name})")

    pages = _EXTRACTORS[ext](path)
    n_pages = len(pages)

    repeated = detect_repeated_lines(pages) if ext == ".pdf" else set()
    raw = _strip_headers_footers(pages, repeated) if ext == ".pdf" else "\n\n".join(pages)

    cleaned = clean_text(raw, drop_references=drop_references)
    return {
        "path": str(path),
        "name": path.name,
        "ext": ext,
        "n_pages": n_pages,
        "n_chars_raw": sum(len(p) for p in pages),
        "n_chars_clean": len(cleaned),
        "removed_headers": sorted(repeated),
        "text": cleaned,
    }
